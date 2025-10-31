"""
Export functionality for DOCX, PDF, and Markdown formats
"""
import asyncio
import json
from typing import Dict, Any, List
from io import BytesIO
import markdown


# For DOCX export
try:
    from docx import Document
    from docx.shared import Inches
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# For PDF export
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# For Markdown export, we'll use built-in functionality


class Exporter:
    """Handle export of redlined documents in various formats"""
    
    def __init__(self, run_data: Dict[str, Any]):
        self.run_data = run_data
    
    async def export_redline(self, format_type: str = "md") -> bytes:
        """Export redlined document in the specified format"""
        if format_type.lower() == "docx":
            if not DOCX_AVAILABLE:
                raise ImportError("python-docx is required for DOCX export. Install with: pip install python-docx")
            return await self._export_docx()
        elif format_type.lower() == "pdf":
            if not PDF_AVAILABLE:
                raise ImportError("reportlab is required for PDF export. Install with: pip install reportlab")
            return await self._export_pdf()
        elif format_type.lower() == "md":
            return await self._export_md()
        else:
            raise ValueError(f"Unsupported format: {format_type}. Supported formats: docx, pdf, md")
    
    async def _export_docx(self) -> bytes:
        """Export as DOCX"""
        doc = Document()
        
        # Add title
        doc.add_heading(f'Contract Redline Report - Run {self.run_data.get("run_id", "Unknown")}', 0)
        
        # Add metadata
        doc.add_paragraph(f"Document ID: {self.run_data.get('doc_id', 'N/A')}")
        doc.add_paragraph(f"Agent Path: {self.run_data.get('agent_path', 'N/A')}")
        doc.add_paragraph(f"Score: {self.run_data.get('score', 0)}")
        doc.add_paragraph("")
        
        # Add assessments
        doc.add_heading('Risk Assessments', level=1)
        
        for assessment in self.run_data.get("assessments", []):
            clause_id = assessment.get("clause_id", "N/A")
            risk_level = assessment.get("risk_level", "N/A")
            rationale = assessment.get("rationale", "N/A")
            
            # Add clause info
            doc.add_heading(f'Clause: {clause_id} (Risk: {risk_level})', level=2)
            doc.add_paragraph(f'Rationale: {rationale}')
            
            # Add policy references if any
            policy_refs = assessment.get("policy_refs", [])
            if policy_refs:
                doc.add_paragraph(f'Policy References: {", ".join(policy_refs)}')
            
            doc.add_paragraph("")
        
        # Add proposals
        doc.add_heading('Redline Proposals', level=1)
        
        for proposal in self.run_data.get("proposals", []):
            clause_id = proposal.get("clause_id", "N/A")
            original_text = proposal.get("original_text", "")
            edited_text = proposal.get("edited_text", "")
            rationale = proposal.get("rationale", "")
            
            # Add proposal info
            doc.add_heading(f'Proposal for Clause: {clause_id}', level=2)
            doc.add_paragraph(f'Original: {original_text}')
            doc.add_paragraph(f'Edited: {edited_text}')
            doc.add_paragraph(f'Rationale: {rationale}')
            
            doc.add_paragraph("")
        
        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer.read()
    
    async def _export_pdf(self) -> bytes:
        """Export as PDF"""
        buffer = BytesIO()
        
        # Create document
        pdf_doc = SimpleDocTemplate(buffer, pagesize=letter)
        elements = []
        styles = getSampleStyleSheet()
        
        # Add title
        title = Paragraph(f'Contract Redline Report - Run {self.run_data.get("run_id", "Unknown")}', styles['Title'])
        elements.append(title)
        elements.append(Spacer(1, 12))
        
        # Add metadata
        metadata_text = f"""
        Document ID: {self.run_data.get('doc_id', 'N/A')}<br/>
        Agent Path: {self.run_data.get('agent_path', 'N/A')}<br/>
        Score: {self.run_data.get('score', 0)}<br/>
        """
        metadata_para = Paragraph(metadata_text, styles['Normal'])
        elements.append(metadata_para)
        elements.append(Spacer(1, 12))
        
        # Add assessments
        assessments_header = Paragraph('Risk Assessments', styles['Heading1'])
        elements.append(assessments_header)
        
        for assessment in self.run_data.get("assessments", []):
            clause_id = assessment.get("clause_id", "N/A")
            risk_level = assessment.get("risk_level", "N/A")
            rationale = assessment.get("rationale", "N/A")
            
            # Clause header
            clause_header_text = f'Clause: {clause_id} (Risk: {risk_level})'
            clause_header = Paragraph(clause_header_text, styles['Heading2'])
            elements.append(clause_header)
            
            # Rationale
            rationale_para = Paragraph(f'Rationale: {rationale}', styles['Normal'])
            elements.append(rationale_para)
            
            # Policy references
            policy_refs = assessment.get("policy_refs", [])
            if policy_refs:
                refs_para = Paragraph(f'Policy References: {", ".join(policy_refs)}', styles['Normal'])
                elements.append(refs_para)
            
            elements.append(Spacer(1, 6))
        
        # Add proposals
        proposals_header = Paragraph('Redline Proposals', styles['Heading1'])
        elements.append(proposals_header)
        
        for proposal in self.run_data.get("proposals", []):
            clause_id = proposal.get("clause_id", "N/A")
            original_text = proposal.get("original_text", "")
            edited_text = proposal.get("edited_text", "")
            rationale = proposal.get("rationale", "")
            
            # Proposal header
            proposal_header_text = f'Proposal for Clause: {clause_id}'
            proposal_header = Paragraph(proposal_header_text, styles['Heading2'])
            elements.append(proposal_header)
            
            # Original text
            original_para = Paragraph(f'Original: {original_text}', styles['Normal'])
            elements.append(original_para)
            
            # Edited text
            edited_para = Paragraph(f'Edited: {edited_text}', styles['Normal'])
            elements.append(edited_para)
            
            # Rationale
            rationale_para = Paragraph(f'Rationale: {rationale}', styles['Normal'])
            elements.append(rationale_para)
            
            elements.append(Spacer(1, 6))
        
        # Build PDF
        pdf_doc.build(elements)
        buffer.seek(0)
        return buffer.read()
    
    async def _export_md(self) -> bytes:
        """Export as Markdown"""
        md_content = f"# Contract Redline Report - Run {self.run_data.get('run_id', 'Unknown')}\n\n"
        
        # Add metadata
        md_content += f"**Document ID:** {self.run_data.get('doc_id', 'N/A')}\n"
        md_content += f"**Agent Path:** {self.run_data.get('agent_path', 'N/A')}\n"
        md_content += f"**Score:** {self.run_data.get('score', 0)}\n\n"
        
        # Add assessments
        md_content += "## Risk Assessments\n\n"
        
        for assessment in self.run_data.get("assessments", []):
            clause_id = assessment.get("clause_id", "N/A")
            risk_level = assessment.get("risk_level", "N/A")
            rationale = assessment.get("rationale", "N/A")
            
            md_content += f"### Clause: {clause_id} (Risk: {risk_level})\n"
            md_content += f"**Rationale:** {rationale}\n"
            
            policy_refs = assessment.get("policy_refs", [])
            if policy_refs:
                md_content += f"**Policy References:** {', '.join(policy_refs)}\n"
            
            md_content += "\n"
        
        # Add proposals
        md_content += "## Redline Proposals\n\n"
        
        for proposal in self.run_data.get("proposals", []):
            clause_id = proposal.get("clause_id", "N/A")
            original_text = proposal.get("original_text", "")
            edited_text = proposal.get("edited_text", "")
            rationale = proposal.get("rationale", "")
            
            md_content += f"### Proposal for Clause: {clause_id}\n"
            md_content += f"**Original:** {original_text}\n"
            md_content += f"**Edited:** {edited_text}\n"
            md_content += f"**Rationale:** {rationale}\n\n"
        
        return md_content.encode('utf-8')


async def export_redline_document(run_data: Dict[str, Any], format_type: str = "md") -> bytes:
    """Export a redlined document in the specified format"""
    exporter = Exporter(run_data)
    return await exporter.export_redline(format_type)